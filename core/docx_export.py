"""Word-Bericht-Generierung mit python-docx."""
import os
import logging
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import config

logger = logging.getLogger(__name__)


def generiere_bericht(
    gebaeude_name: str,
    gebaeude_adresse: str | None,
    von_datum: date,
    bis_datum: date,
    zaehler_daten: list,
    ersteller: str = "",
) -> str:
    """
    Generiert einen Word-Bericht für ein Gebäude.
    
    Args:
        gebaeude_name: Name des Gebäudes
        gebaeude_adresse: Adresse (optional)
        von_datum: Beginn des Berichtszeitraums
        bis_datum: Ende des Berichtszeitraums
        zaehler_daten: Liste mit Zähler-Dicts (id, typ, info, ablesungen, verbrauch_zeitraum)
        ersteller: Name des Erstellers
    
    Returns:
        Pfad zur generierten DOCX-Datei
    """
    doc = Document()
    
    # ─── Titel ───────────────────────────────────────────────────────────
    titel = doc.add_heading("Zählerbericht", level=0)
    titel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Untertitel
    untertitel = doc.add_paragraph()
    untertitel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = untertitel.add_run(gebaeude_name)
    run.bold = True
    run.font.size = Pt(16)
    
    if gebaeude_adresse:
        doc.add_paragraph(gebaeude_adresse).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Zeitraum
    zeitraum = doc.add_paragraph()
    zeitraum.alignment = WD_ALIGN_PARAGRAPH.CENTER
    zeitraum.add_run(
        f"Zeitraum: {von_datum.strftime('%d.%m.%Y')} – {bis_datum.strftime('%d.%m.%Y')}"
    )
    
    doc.add_paragraph()  # Abstand
    
    # ─── Zusammenfassung ─────────────────────────────────────────────────
    doc.add_heading("Zusammenfassung", level=1)
    
    # Statistiken berechnen
    gesamt_zaehler = len(zaehler_daten)
    gesamt_ablesungen = sum(len(z["ablesungen"]) for z in zaehler_daten)
    
    # Verbrauch nach Typ gruppieren
    verbrauch_nach_typ = {}
    for z in zaehler_daten:
        typ = z["typ"]
        info = z["info"]
        if z["verbrauch_zeitraum"] is not None:
            if typ not in verbrauch_nach_typ:
                verbrauch_nach_typ[typ] = {
                    "name": info["name"],
                    "einheit": info["einheit"],
                    "icon": info["icon"],
                    "verbrauch": 0,
                }
            verbrauch_nach_typ[typ]["verbrauch"] += z["verbrauch_zeitraum"]
    
    # Zusammenfassung als Tabelle
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Kennzahl"
    hdr_cells[1].text = "Wert"
    
    # Header fett
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Zeilen hinzufügen
    row = table.add_row().cells
    row[0].text = "Anzahl Zähler"
    row[1].text = str(gesamt_zaehler)
    
    row = table.add_row().cells
    row[0].text = "Anzahl Ablesungen"
    row[1].text = str(gesamt_ablesungen)
    
    for typ, daten in verbrauch_nach_typ.items():
        row = table.add_row().cells
        row[0].text = f"Verbrauch {daten['name']}"
        row[1].text = f"{daten['verbrauch']:,.1f} {daten['einheit']}".replace(",", "X").replace(".", ",").replace("X", ".")
    
    doc.add_paragraph()
    
    # ─── Zählerübersicht ─────────────────────────────────────────────────
    doc.add_heading("Zählerübersicht", level=1)
    
    # Tabelle mit allen Zählern
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    
    hdr_cells = table.rows[0].cells
    headers = ["Zähler", "Typ", "Standort", "Aktueller Stand", "Verbrauch"]
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    for z in zaehler_daten:
        info = z["info"]
        row = table.add_row().cells
        row[0].text = f"#{z['id']}"
        row[1].text = f"{info['icon']} {info['name']}"
        row[2].text = z["standort"] or "–"
        
        if z["aktueller_stand"] is not None:
            row[3].text = f"{z['aktueller_stand']:,.1f} {info['einheit']}".replace(",", "X").replace(".", ",").replace("X", ".")
        else:
            row[3].text = "–"
        
        if z["verbrauch_zeitraum"] is not None:
            row[4].text = f"{z['verbrauch_zeitraum']:,.1f} {info['einheit']}".replace(",", "X").replace(".", ",").replace("X", ".")
        else:
            row[4].text = "–"
    
    doc.add_paragraph()
    
    # ─── Detaillierte Ablesungen ─────────────────────────────────────────
    doc.add_heading("Ablesungen im Detail", level=1)
    
    for z in zaehler_daten:
        info = z["info"]
        
        # Zähler-Überschrift
        heading = f"{info['icon']} Zähler #{z['id']} – {info['name']}"
        if z["standort"]:
            heading += f" ({z['standort']})"
        doc.add_heading(heading, level=2)
        
        if z["zaehlernummer"]:
            doc.add_paragraph(f"Zählernummer: {z['zaehlernummer']}")
        
        if not z["ablesungen"]:
            doc.add_paragraph("Keine Ablesungen im Berichtszeitraum.")
            continue
        
        # Ablesungs-Tabelle
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Datum"
        hdr_cells[1].text = f"Stand ({info['einheit']})"
        hdr_cells[2].text = f"Verbrauch ({info['einheit']})"
        
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for a in z["ablesungen"]:
            row = table.add_row().cells
            row[0].text = a["datum"].strftime("%d.%m.%Y")
            row[1].text = f"{a['stand']:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".")
            row[2].text = f"{a['verbrauch']:,.1f}".replace(",", "X").replace(".", ",").replace("X", ".") if a["verbrauch"] else "–"
        
        doc.add_paragraph()
    
    # ─── Footer ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run(f"Erstellt am {date.today().strftime('%d.%m.%Y')}")
    if ersteller:
        footer.add_run(f" von {ersteller}")
    
    # ─── Speichern ───────────────────────────────────────────────────────
    os.makedirs(config.OUTPUT_DIR, exist_ok=True)
    
    safe_name = gebaeude_name.replace(" ", "_").replace("/", "-")
    dateiname = f"Zaehlerbericht_{safe_name}_{von_datum.strftime('%Y-%m')}.docx"
    dateipfad = os.path.join(config.OUTPUT_DIR, dateiname)
    
    doc.save(dateipfad)
    logger.info(f"Bericht gespeichert: {dateipfad}")
    
    return dateipfad
